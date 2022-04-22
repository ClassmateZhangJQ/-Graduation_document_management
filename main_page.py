import time
import numpy as np
# from memory_pic import ahpu_png,artistic_body01_png,owners_png
import tkinter as tk
import ttkbootstrap as ttk
from tkinter import messagebox
from PIL import  Image, ImageTk
from tkinter import filedialog
from docx import Document
import pandas as pd
from PIL import Image, ImageTk
from threading import Thread
from multiprocessing import Process
import os
import re
from win32com import client as wc #导入模块
import shutil
from Title_sims import Title_sims
from win32com.client import constants,gencache
import threading
from concurrent.futures import ThreadPoolExecutor
class MainPage():

    def __init__(self, master: tk.Tk):
        self.root = master
        self.width = 850
        self.height = 650
        self.x = (self.root.winfo_screenwidth() - self.width) / 2
        self.y = (self.root.winfo_screenheight() - self.height) / 2
        self.root.title("毕业文档批量处理小助手 1.1.1.220418_alpha")
        self.root.geometry("%dx%d+%d+%d" % (self.width,self.height,self.x,self.y))
        self.root.resizable(False,False)
        self.root.protocol('WM_DELETE_WINDOW',self.exit_editor)


        self.tv_tables1 = ['编号','题目名称','学生姓名','组别']
        self.tv_tables2 = ['文档名称','字数','类别','是否达标']
        self.tv_tables3 = ['文档名称',]
        self.tv_tables4 = ['文档名称',]
        self.tv_table = ['操作时间','操作功能','状态']


        self.artistic_body = Image.open('./logo/artistic_body01.png').resize((575,85))
        self.artistic_body = ImageTk.PhotoImage(self.artistic_body)

        self.ahpu = Image.open('./logo/ahpu.png').resize((280,85))
        self.ahpu = ImageTk.PhotoImage(self.ahpu)
        self.excel_list = []
        ############ >>>>>>>>>>>>>>>>>> 根据审批表生成word文档
        self.fuc01_1 = ttk.StringVar() # 审批表excel 路径
        self.fuc01_2 = ttk.StringVar() # 上传停用词
        self.fuc01_3 = ttk.StringVar() # 审批表功能内的保存路径
        self.fuc01_4_1 =ttk.StringVar()
        self.fuc01_4_2 =ttk.StringVar()
        ############ >>>>>>>>>>>>>>>>>> 成绩评定表批量插入签名
        self.fuc09_1 = ttk.StringVar() # 答辩表所在的目录
        self.fuc09_2 = ttk.StringVar() # 插入签名所在的位置
        self.fuc09_3 = ttk.StringVar() # combobox
        self.fuc09_4 = ttk.StringVar() # data
        ############ >>>>>>>>>>>>>>>>>> 根据选题表excel批量生成对应\n    成绩评定表、答辩记录表
        self.fuc06_1 = ttk.StringVar()
        self.fuc06_2 = ttk.StringVar()
        self.fuc06_3 = ttk.StringVar()
        self.fuc06_4 = ttk.StringVar() # 日期
        self.fuc06_5 = ttk.StringVar() # 届数
        self.fuc06_6 = ttk.StringVar() # 模板
        self.fuc06_7 = ttk.StringVar() # 学院
        self.fuc06_8 = ttk.StringVar() # 保存格式
        self.field6 = ['可选字段名']
        ############ >>>>>>>>>>>>>>>>>> 开题报告，任务书批量插入签名
        self.fuc05_1 = ttk.StringVar() # 模板保存的路径
        self.fuc05_2 = ttk.StringVar() # commbox word类型
        self.fuc05_3 = ttk.StringVar() # 表所在的目录
        self.fuc05_4 = ttk.StringVar() # commbox 身份id
        self.fuc05_5 = ttk.StringVar() # 签名

        ############ >>>>>>>>>>>>>>>>>> # 字数检查
        self.fuc03_1 = ttk.StringVar() # 表所在目录
        self.fuc03_2 = ttk.IntVar() # 字数
        self.fuc03_3 = ttk.StringVar() # 身份

        ############ >>>>>>>>>>>>>>>>>> # 合并word文档
        self.fuc07_1 = ttk.StringVar()  # 目录
        self.fuc07_2 = ttk.StringVar()  # 合并后的名称
        self.fuc07_3 = ttk.StringVar()

        ############ >>>>>>>>>>>>>>>>>  # 批量产生pdf
        self.fuc08_1 = ttk.StringVar()

        self.pool = ThreadPoolExecutor(max_workers=3)


        self.tree_view_data1 = []
        self.tree_stui1 = 1

        self.tree_view_data2 = []
        self.tree_view_data3 = [] # 合并word文档

        self.tree_work_info3 = [] # 合并word功能内存储点击事件内容
        self.tree_work_info4 = [] # 批量生成pdf功能
        self.tree_view_logger = []


        self.create_page()

    def exit_editor(self):
        if messagebox.askokcancel('退出?', '确定退出吗?'):  # 设置文本提示框
            self.root.destroy()

    def create_page(self):
        self.lowerest_frame = ttk.Frame(self.root,width = self.width,height =self.height,bootstyle = 'light').place(x=0,y=0)
        self.artistic_pho = tk.Label(self.root,image = self.artistic_body,width = 575,).place(rely = 0 , relx = 0.3345)
        self.ahpu_pho = tk.Label(self.root,image = self.ahpu).place(relx = 0,rely = 0)
        self.onwers_logo_name = ttk.Label(self.root,text = '出品：软件211章佳乾  |  数科211张长峰  |  数据科学211 鲁嘉琪 ©安徽工程大学BDAI协会',width = 200,bootstyle = 'dark').place(rely = 0.96,relx = 0.34)
        self.left_frame = ttk.Frame(self.root,width = self.width/5+40,height = 5 *self.height/6-60).place(rely = 0.18,relx = 0)


        ttk.Label(self.left_frame,text = '                 审批阶段                    ',bootstyle = 'inverse-danger').place(relx = 0.01, rely = 0.15)
        ttk.Label(self.left_frame,text = '               中期检查阶段                ',bootstyle = 'inverse-danger').place(relx = 0.01, rely = 0.375)
        ttk.Label(self.left_frame,text = '                 答辩阶段                    ',bootstyle = 'inverse-danger').place(relx = 0.01, rely = 0.575)

        ttk.Button(self.left_frame,text = 'excel选题表\n  辅助查重',width = 25,command = self.show_fuc1,bootstyle = 'light').place(relx = 0.01, rely = 0.184)

        ttk.Label(self.left_frame,text = '意见反馈：qq群752885480',bootstyle = 'dark').place(rely = 0.96,relx = 0.02)
        ttk.Button(self.left_frame,text = '开题报告、任务书\n  批量插入签名',  width = 25,command = self.show_fuc5,bootstyle = 'light').place(relx=0.01, rely=0.409)

        ttk.Button(self.left_frame,text = '根据选题表excel批量生成对应\n    成绩评定表、答辩记录表', width = 25,bootstyle = 'light',command = self.show_fuc6).place(relx=0.01, rely = 0.685)
        ttk.Button(self.left_frame,text = '合并word文档',  width = 25,bootstyle = 'light', command = self.show_fuc7).place(relx=0.01, rely = 0.830)
        ttk.Button(self.left_frame,text = '批量生成pdf',  width = 25,bootstyle = 'light',command = self.show_fuc8 ).place(relx=0.01, rely = 0.875)
        ttk.Button(self.left_frame,text = '   成绩评定表\n教师评语字数检查',  width = 25, command = self.show_fuc3,bootstyle = 'light').place(relx=0.01, rely=0.758)
        ttk.Button(self.left_frame,text = '成绩评定表、答辩记录表\n         批量插入签名',width = 25,bootstyle = 'light',command = self.show_fuc9).place(relx=0.01, rely = 0.609)
        self.bootom_frame = ttk.LabelFrame(self.root,text = '日志',width = 5 * self.width /6-95,height = self.height /4,bootstyle ='dark').place(rely = 0.67,relx = 0.25)
        self.excel_related_f1_frame = ttk.LabelFrame(self.root,text = 'excel选题表辅助查重',width = 5 * self.width /6-95,height = 4 * self.height /6 -85,bootstyle ='dark')  #  excel选题表批量查重
        self.insert_font_f1_frame = ttk.LabelFrame(self.root,text = '开题报告\任务书\n批量插入签名',width = 5 * self.width /6-95,height = 4 * self.height /6 -85,bootstyle ='dark')  #  批量插入签名
        self.check_font_nums_frame = ttk.LabelFrame(self.root,text = '检查字数',width = 5 * self.width /6-95,height = 4 * self.height /6 -85,bootstyle ='dark')  #  检查字数
        self.task_word_f1frame = ttk.LabelFrame(self.root,text = '开题报告、任务书批量插入签名',width = 5 * self.width /6-95,height = 4 * self.height /6 -85,bootstyle ='dark')  #  任务书批量插入签名
        self.finall_check_scores = ttk.LabelFrame(self.root,text = '根据选题表excel批量生成答辩记录表、成绩评定表',width = 5 * self.width /6-95,height = 4 * self.height /6 -85,bootstyle ='dark')  #  成绩评定表检查字数
        self.combine_words = ttk.LabelFrame(self.root,text = '合并word文档',width = 5 * self.width /6-95,height = 4 * self.height /6 -85,bootstyle ='dark')  #  合并文档
        self.auto_make_pdf = ttk.LabelFrame(self.root,text = '批量生成pdf',width = 5 * self.width /6-95,height = 4 * self.height /6 -85,bootstyle ='dark')  #  批量生成pdf
        self.auto_insert_name = ttk.LabelFrame(self.root,text = '成绩评定表、答辩记录表批量插入签名',width = 5 * self.width /6-95,height = 4 * self.height /6 -85,bootstyle ='dark')  #  批量生成pdf

        self.logger = ttk.Treeview(self.bootom_frame,show = "headings",columns =self.tv_table)
        self.logger.place(relx =0.26,rely = 0.7,width = 595,height = 130)
        for i in range(len(self.tv_table)):
            self.logger.heading(column=self.tv_table[i], text=self.tv_table[i], anchor='center',)  # 设置除#0列以外的表头
            self.logger.column(self.tv_table[i], width=25, anchor='center', stretch=True)

        self.progressbar9 = ttk.Progressbar(self.auto_insert_name,bootstyle = 'success')
        self.progressbar8 = ttk.Progressbar(self.auto_make_pdf,bootstyle = 'success')
        self.progressbar7 = ttk.Progressbar(self.combine_words,bootstyle = 'success')
        self.progressbar6 = ttk.Progressbar(self.finall_check_scores,bootstyle = 'success')
        self.progressbar5 = ttk.Progressbar(self.task_word_f1frame,bootstyle = 'success')
        self.progressbar3 = ttk.Progressbar(self.check_font_nums_frame,bootstyle = 'success')


    def show_fuc1(self): #  excel选题表批量查重
        self.insert_font_f1_frame.place_forget()
        self.check_font_nums_frame.place_forget()
        self.task_word_f1frame.place_forget()
        self.auto_insert_name.place_forget()
        self.finall_check_scores.place_forget()
        self.combine_words.place_forget()
        self.auto_make_pdf.place_forget()

        self.excel_related_f1_frame.place(rely = 0.133,relx = 0.25)
        yscroll = ttk.Scrollbar(self.root,orient = 'vertical')
        ttk.Label(self.excel_related_f1_frame,text = '使用说明:\n1.请上传选题表(excel文件)和停用词(txt文件)\n2.点击开始按钮(本功能将相似的题目定为一组,方便人工二次筛查甄别)').place(rely = 0.82,relx = 0)
        ttk.Button(self.excel_related_f1_frame,text = '点击上传选题表(Excel)',width = 23,bootstyle="outline-dark",command = lambda:self.get_fuc_excel(self.fuc01_1)).place(relx = 0.01,rely = 0.01)
        ttk.Entry(self.excel_related_f1_frame,textvariable = self.fuc01_1,width = 31,state="readonly").place(relx = 0.325,rely = 0.01)
        ttk.Button(self.excel_related_f1_frame,text = '点击上传停用词',bootstyle="outline-dark",command = lambda :self.get_fuc_model(self.fuc01_2,'上传停用词')).place(rely = 0.01,relx = 0.73)
        self.tv1 = ttk.Treeview(self.excel_related_f1_frame,show = "headings",columns =self.tv_tables1,yscrollcommand = yscroll.set,bootstyle='light')
        self.tv1.place(rely = 0.105,relx = 0.01,width = 600,height = 200)
        self.tv1_flag = 1
        yscroll.config(command = self.tv1.yview)

        ttk.Button(self.excel_related_f1_frame,text = '点击开始查重',command = lambda:self.pool.submit(self.run_fuc,self.fuc1_run),bootstyle="outline-dark").place(rely = 0.72, relx = 0.01)
        self.progressbar1 = ttk.Progressbar(self.excel_related_f1_frame,bootstyle = 'success')
        self.progressbar1.place(rely= 0.725,relx = 0.185,width = 185,height = 20)

        for i in range(len(self.tv_tables1)):
            self.tv1.heading(column=self.tv_tables1[i], text=self.tv_tables1[i], anchor='center',)  # 设置除#0列以外的表头
            self.tv1.column(self.tv_tables1[i], width=50, anchor='center')
        self.tree_stui1 = 0

        if len(self.tree_view_data1) and self.tv1_flag:
            self.insert2treeview(tree_view_data=self.tree_view_data1,tv=self.tv1)
            self.tv1_flag = 0
    def fuc1_run(self):
        now = time.strftime("%Y-%m-%d %X",time.localtime())
        self.tree_view_data1 = []
        excel_data = pd.read_excel(self.fuc01_1.get())
        height , width = excel_data.shape
        data = []
        self.progressbar1['maximum'] = height
        self.progressbar1['value'] = 0
        stopwords = [line.strip() for line in open(self.fuc01_2.get(), encoding='utf-8').readlines()]
        for i in range(1,height):
            data.append(excel_data.iloc[i][1].replace('\n','').replace('\t',''))
        ret = Title_sims(data,0.5,height-1,stopwords).ret
        key_list = [0, 1, 8]
        class_num = 0
        for per_class in ret:
            for title_ in per_class:
                for i in range(1, height):
                    flag = title_ == (excel_data.iloc[i][key_list[1]].replace('\n', '').replace('\t', ''))
                    if flag:
                        self.tree_view_data1.append([excel_data.iloc[i][key_list[0]],excel_data.iloc[i][key_list[1]],excel_data.iloc[i][key_list[2]],class_num])
                        break
                self.progressbar1['value'] += len(per_class)
            class_num += 1
        self.progressbar1['value'] = height
        if self.tv1_flag:
            for child in self.tree_view_data1:
                self.tv1.insert('', 'end', values=child)
            self.tv_flag = 0
        self.logger.insert('','end',value = [now,'选题表查重','正常'] )

    def show_fuc3(self):  #  检查字数
        self.excel_related_f1_frame.place_forget()
        self.insert_font_f1_frame.place_forget()
        self.task_word_f1frame.place_forget()
        self.auto_insert_name.place_forget(),
        self.finall_check_scores.place_forget()
        self.combine_words.place_forget()
        self.auto_make_pdf.place_forget()
        self.check_font_nums_frame.place(rely = 0.133,relx = 0.25)
        ttk.Label(self.check_font_nums_frame,text = '使用说明:\n1.选择需要处理的的成绩评定表所在的目录(注意:该目录下只能存在成绩评定表)\n2.输入数字数量下限后，点击开始，等待进度条结束').place(rely = 0.82,relx = 0.02)
        ttk.Button(self.check_font_nums_frame,text = '点击选择成绩评定表所在目录',bootstyle="outline-dark",width = 23,command = lambda:self.get_fuc_word_menu(self.fuc03_1)).place(relx = 0.02,rely = 0.02)
        ttk.Entry(self.check_font_nums_frame,textvariable = self.fuc03_1,width = 31,state="readonly").place(relx = 0.325,rely = 0.02)
        ttk.Label(self.check_font_nums_frame,text = '输入字数数量下限：').place(relx=0.022, rely=0.145)
        ttk.Entry(self.check_font_nums_frame,textvariable = self.fuc03_2,width = 4).place(relx = 0.21,rely = 0.127)
        ttk.Button(self.check_font_nums_frame,text = '开始',bootstyle="outline-dark",command = lambda:self.pool.submit(self.run_fuc,self.fuc3_run),width = 10).place(relx = 0.02, rely = 0.72)
        ttk.Combobox(self.check_font_nums_frame,value = ['指导教师评语','评阅教师评语','答辩小组评语'],bootstyle="secondary",textvariable = self.fuc03_3).place(relx = 0.325,rely = 0.127)
        yscroll = ttk.Scrollbar(self.root,orient = 'vertical')
        self.tv2 = ttk.Treeview(self.check_font_nums_frame,show = "headings",columns =self.tv_tables2,height=30,yscrollcommand = yscroll.set)
        self.tv2.place(rely = 0.245,relx = 0.02,width = 570,height = 150)
        self.tv2_flag = 1
        yscroll.config(command = self.tv2.yview)
        for i in range(len(self.tv_tables2)):
            self.tv2.heading(column=self.tv_tables2[i], text=self.tv_tables2[i], anchor='center',
                               command=lambda: print(self.tv_tables2[i]))  # 设置除#0列以外的表头
            self.tv2.column(self.tv_tables2[i], width=25, anchor='center', stretch=True)
        self.progressbar3.place(rely= 0.732,relx = 0.185,width = 185,height = 20)
        if len(self.tree_view_data2) and self.tv2_flag:
            self.insert2treeview(tree_view_data=self.tree_view_data2,tv=self.tv2)

    def fuc3_run(self):
        now = time.strftime("%Y-%m-%d %X",time.localtime())
        self.tree_view_data2 = []
        word_menu_path = self.fuc03_1.get()
        word_list = os.listdir(word_menu_path)
        limit_num = self.fuc03_2.get()
        id_class = self.fuc03_3.get()
        for word in word_list:
            word_path = os.path.join(word_menu_path,word)
            doc = Document(word_path)
            height = len(doc.tables[0].rows)
            flag = True
            for h in range(height):
                list_ = doc.tables[0].rows[h].cells
                if len(list_) == len(list(set(list_))) or not list_[0].text.count(id_class):
                    continue #  到下一行
                cell_txt = list_[0].text
                ret = re.findall('(.*?：)', cell_txt)
                if len(ret) == 0:
                    ret = re.findall('(.*?:)', cell_txt)
                fir = ret[0]
                fir_index = cell_txt.find(fir) + len(fir) + 1
                sec = ret[1]
                sec_index = cell_txt.find(sec)
                final_text = cell_txt[fir_index:sec_index]
                count_ = final_text.count('\n') # 统计\n 的数量
                ret_count = len(final_text) - count_
                if ret_count < limit_num:
                    flag = False
                self.tree_view_data2.append([word,ret_count,id_class,flag])
        if self.tv2_flag:
            for child in self.tree_view_data2:
                self.tv2.insert('', 'end', values=child)
            self.tv2_flag = 0
        self.logger.insert('','end',value = [now,'成绩评语字数检查','正常'] )


    def show_fuc5(self):  # 任务书批量插入签名
        self.excel_related_f1_frame.place_forget()
        self.insert_font_f1_frame.place_forget()
        self.check_font_nums_frame.place_forget()
        self.auto_insert_name.place_forget()
        self.finall_check_scores.place_forget()
        self.combine_words.place_forget()
        self.auto_make_pdf.place_forget()
        self.task_word_f1frame.place(rely = 0.133,relx = 0.25)

        ttk.Label(self.task_word_f1frame,text = '选择需要处理的word类型：').place(relx = 0.02,rely = 0.02)
        ttk.Combobox(self.task_word_f1frame,value = ['任务书','开题报告'],bootstyle="secondary",textvariable = self.fuc05_2).place(relx = 0.325,rely = 0.02)
        # ttk.Button(self.task_word_f1frame,text = '点击上传模板',width = 15,bootstyle = "outline-dark",command =lambda: self.get_fuc_model(self.fuc05_1)).place(rely = 0.02,relx = 0.6025)
        ttk.Button(self.task_word_f1frame,text = '点击选择表所在目录',width = 23,bootstyle="outline-dark",command = lambda:self.get_fuc_word_menu(self.fuc05_3)).place(relx=0.017, rely=0.122)
        ttk.Entry(self.task_word_f1frame,textvariable = self.fuc05_3,width = 31,state="readonly").place(relx = 0.325,rely = 0.122)
        ttk.Label(self.task_word_f1frame,text = '选择您的身份：').place(rely = 0.25,relx = 0.02)
        ttk.Combobox(self.task_word_f1frame,value = ['指导教师','教研室主任'],bootstyle="secondary",textvariable = self.fuc05_4).place( rely =  0.25,relx = 0.17,width = 125,height = 30)
        ttk.Button(self.task_word_f1frame,text = '录入签名图片',width = 10,bootstyle="outline-dark",command = lambda : self.get_fuc_label(self.fuc05_5)).place(relx=0.40, rely= 0.25)
        ttk.Label(self.task_word_f1frame,text = '使用说明:\n1.选择需要批量插入签名的任务书所在的目录(注意:该目录下只能存在任务书word)\n2.录入签名后点击开始，等待进度条结束').place(rely = 0.82,relx = 0.02)
        ttk.Button(self.task_word_f1frame,text = '开始',bootstyle="outline-dark",command =lambda:self.pool.submit(self.run_fuc,self.fuc5_run),width = 10).place(relx = 0.02, rely = 0.72)
        self.progressbar5.place(rely= 0.732,relx = 0.185,width = 185,height = 20)

    def fuc5_run(self):
        now = time.strftime("%Y-%m-%d %X",time.localtime())
        word_menu = self.fuc05_3.get()
        my_id = self.fuc05_4.get()
        word_class = self.fuc05_2.get()
        pic_path = self.fuc05_5.get()
        # doc = Document(model_path)
        word_list = os.listdir(word_menu)
        self.progressbar5['maximum'] = len(word_list)
        self.progressbar5['value'] = 0
        for i in range(len(word_list)):
            path = word_menu + '\\' + word_list[i]
            doc = Document(path)
            if word_class == '开题报告':
                image = self.change_pic_size(pic_path, (15, 15))
                new_path = pic_path[:-4] + 'new.png'
                image.save(new_path)
                for i in doc.paragraphs:
                    if i.text.count(my_id):
                        run = i.add_run()
                        run.add_picture(new_path)
            elif word_class == '任务书':
                if my_id == '教研室主任':
                    return
                image = self.change_pic_size(pic_path, (15, 15))
                new_path = pic_path[:-4] + 'new.png'
                image.save(new_path)
                for i in doc.paragraphs:
                    if i.text.count(my_id):
                        run = i.add_run()
                        run.add_picture(new_path)
            self.progressbar5['value'] += 1
            doc.save(path)
            self.task_word_f1frame.update()
        self.logger.insert('','end',value = [now,'开题报告\\任务书插入签名','正常'] )


    def show_fuc6(self):  #  答辩记录表操作
        self.excel_related_f1_frame.place_forget()
        self.insert_font_f1_frame.place_forget()
        self.check_font_nums_frame.place_forget()
        self.auto_insert_name.place_forget()
        self.combine_words.place_forget()
        self.auto_make_pdf.place_forget()
        self.task_word_f1frame.place_forget()
        self.fuc06_4.set('')
        self.finall_check_scores.place(rely = 0.133,relx = 0.25)

        ttk.Label(self.finall_check_scores,text = '使用说明：\n1.选择word类型后，上传相应模板\n2.点击上传选题表,选择保存路径和做好相应配置后点击开始即可').place(relx =0.02,rely = 0.82)
        ttk.Button(self.finall_check_scores,text = '点击选择选题表(excel)',width = 23,bootstyle="outline-dark",command = lambda:self.fuc6_excel_realte2combox()).place(relx=0.017, rely=0.122)
        ttk.Entry(self.finall_check_scores,textvariable = self.fuc06_1,width = 31,state="readonly").place(relx = 0.325,rely = 0.122)
        ttk.Label(self.finall_check_scores,text = '选择需要处理的word类型：').place(relx = 0.02,rely = 0.02)
        ttk.Combobox(self.finall_check_scores,value = ['成绩评定表','答辩记录表'],bootstyle="secondary",textvariable = self.fuc06_2).place(relx = 0.325,rely = 0.02)

        ttk.Button(self.finall_check_scores,text = '点击选择保存路径',width = 23,bootstyle="outline-dark",command = lambda :self.get_fuc_savepath(self.fuc06_3)).place(relx=0.017, rely=0.235)

        ttk.Button(self.finall_check_scores,text = '点击上传模板',width = 15,bootstyle = "outline-dark",command = lambda :self.get_fuc_model(self.fuc06_6,'上传模板')).place(rely = 0.02,relx = 0.6025)
        ttk.Entry(self.finall_check_scores,textvariable = self.fuc06_3,width = 31,state="readonly").place(relx = 0.325,rely = 0.235)
        ttk.Label(self.finall_check_scores,text = '所属学院:').place(relx = 0.01,rely = 0.455)
        ttk.Combobox(self.finall_check_scores,value = ['计算机与信息学院','艺术学院','体育学院','建筑工程学院','数理与金融学院','人工智能学院','纺织服装学院','化学与环境工程学院','生物与食品工程学院','经济与管理学院','机械工程学院','电气工程学院','材料科学与工程学院',],textvariable = self.fuc06_7).place(relx = 0.12,rely = 0.455)
        self.combobox6 = ttk.Combobox(self.finall_check_scores,value = self.field6,)
        self.combobox6.current(0)
        self.combobox6.place(relx = 0.695,rely = 0.56)

        ttk.Label(self.finall_check_scores,text = '请选择日期:').place(relx = 0.01,rely = 0.36)
        ttk.Label(self.finall_check_scores,text = '请输入当前所属届').place(relx = 0.425,rely = 0.36)
        ttk.Entry(self.finall_check_scores,textvariable = self.fuc06_5,width = 5).place(relx = 0.6,rely = 0.35)
        ttk.Button(self.finall_check_scores,text = '开始',bootstyle="outline-dark",command = lambda:self.pool.submit(self.run_fuc,self.fuc6_run),width = 10).place(relx = 0.02, rely = 0.72)
        ttk.Label(self.finall_check_scores,text = '选择保存文件格式:').place(relx = 0.01,rely = 0.575)
        ttk.Entry(self.finall_check_scores,textvariable = self.fuc06_8,width = 40).place(relx = 0.2,rely = 0.56)
        ttk.Label(self.finall_check_scores,text = '格式说明：{自定义名称}{@字段名}{@字段名}').place(relx = 0.01,rely = 0.654)
        self.fuc6_data = ttk.DateEntry(self.finall_check_scores,bootstyle = 'secondary',textvariable = self.fuc06_4).place(rely = 0.35,relx = 0.12)
        # self.fuc6_text = ttk.ScrolledText(self.finall_check_scores,width = 80,state = 'disabled',highlightcolor=style.colors.primary,highlightbackground = style.colors.border)
        self.progressbar6.place(rely= 0.74,relx = 0.185,width = 185,height = 20)

    def fuc6_run(self):
        now = time.strftime("%Y-%m-%d %X",time.localtime())
        com_class = self.fuc06_2.get()
        sessions = self.fuc06_5.get() # 所属届
        save_path = self.fuc06_3.get()
        bl_school = self.fuc06_7.get()
        model_path = self.fuc06_6.get() # 模板
        excel_data = self.fuc06_1.get()
        the_day = self.fuc06_4.get()
        doc = Document(model_path)
        df = pd.read_excel(excel_data)
        data_arry = np.array(df)
        data_list = data_arry.tolist()
        if len(self.excel_list) == 0:
            self.excel_list = self.get_excel_para(df)
        height,width = df.shape
        test = self.fuc06_8.get()
        obj1 = re.compile('{(?P<a>.*?)}')
        ret1 = obj1.match(test)
        obj2 = re.compile('{@(?P<b>.*?)}')
        ret2 = obj2.findall(test)
        temp_list = [] # 存放index 顺序
        insert_data = the_day.split('/')
        self.progressbar6['maximum'] = height - 1
        print(height)
        self.progressbar6['value'] = 0
        for i in range(len(ret2)):
            temp_list.append(self.excel_list.index(ret2[i]))
        if com_class == '成绩评定表':
            for i in range(1,height):
                path = save_path + '\\' + ret1.group('a')
                doc.tables[0].rows[0].cells[1].text = data_list[i][8]
                doc.tables[0].rows[0].cells[3].text = data_list[i][7]  # 专业班级 7
                doc.tables[0].rows[0].cells[5].text = str(data_list[i][9])  # 学号 9
                doc.tables[0].rows[1].cells[1].text = data_list[i][1]  # 课题名称 1
                doc.paragraphs[1].text = '安徽工程大学{}届本科毕业设计（论文）成绩评定表'.format(sessions)
                doc.paragraphs[3].text = '学院(公章)：{}'.format(bl_school)
                doc.tables[0].rows[5].cells[0].paragraphs[5].text = '日    期：{}年{}月{}日'.format(insert_data[0],insert_data[1],insert_data[2])
                for j in range(len(temp_list)):
                    data_list[i][temp_list[j]].replace('\n', '')
                    data_list[i][temp_list[j]].replace('\t', '')
                    path += data_list[i][temp_list[j]]
                doc.save(path + '.docx')
                self.progressbar6['value'] += 1
                self.finall_check_scores.update()
        else:
            for i in range(1,height):
                path = save_path + '\\' + ret1.group('a')
                doc.tables[0].rows[0].cells[1].text = data_list[i][8]
                doc.tables[0].rows[0].cells[3].text = data_list[i][7]
                doc.tables[0].rows[0].cells[5].text = str(data_list[i][9])
                doc.tables[0].rows[0].cells[7].text = data_list[i][3]
                doc.tables[0].rows[1].cells[1].text = data_list[i][1]
                doc.paragraphs[1].text = '安徽工程大学{}届本科'.format(sessions)
                doc.paragraphs[4].text = '学院名称:{}'.format(bl_school)
                doc.tables[0].rows[16].cells[0].paragraphs[2].text = '{}年{}月{}日'.format(insert_data[0],insert_data[1],insert_data[2])
                for j in range(len(temp_list)):
                    path += data_list[i][temp_list[j]].replace('\n','')
                doc.save(path + '.docx')
                self.progressbar6['value'] += 1
                self.finall_check_scores.update()
        self.logger.insert('','end',value = [now,'批量生成关联文档','正常'] )


    def show_fuc7(self):  #  合并word文档
        self.excel_related_f1_frame.place_forget()
        self.insert_font_f1_frame.place_forget()
        self.check_font_nums_frame.place_forget()
        self.auto_insert_name.place_forget()
        self.auto_make_pdf.place_forget()
        self.task_word_f1frame.place_forget()
        self.finall_check_scores.place_forget()

        self.combine_words.place(rely = 0.133,relx = 0.25)
        ttk.Button(self.combine_words,text =' 点击上传文档所在目录',width = 23,bootstyle="outline-dark",command = lambda :self.get_fuc_word_menu(self.fuc07_1) ).place(relx = 0.02,rely = 0.02)
        ttk.Label(self.combine_words,text = '使用说明:\n1.上传需要被合并的文档所在目录,然后点击显示目录\n2.选中预览图中的需要被合并的文档名(按住shift选中多行)，输入要修改的名称后点击开始合并即可').place(rely = 0.82,relx = 0.02)
        ttk.Entry(self.combine_words,width = 31,state = 'readonly',textvariable = self.fuc07_1).place(relx = 0.325,rely = 0.02)
        ttk.Button(self.combine_words,text = '开始合并',command = lambda:self.pool.submit(self.run_fuc,self.fuc7_run),bootstyle="outline-dark",width = 10).place(relx = 0.02, rely = 0.72)
        ttk.Button(self.combine_words,text = '显示目录',command = self.fuc7_vis_menu,bootstyle="outline-dark",).place(relx = 0.708,rely = 0.02)
        ttk.Label(self.combine_words,text = '请输入选中word合并后的名称:').place(relx=0.02,rely =0.145)
        ttk.Entry(self.combine_words,textvariable = self.fuc07_2,width =15).place(relx=0.325,rely = 0.127)
        self.progressbar7.place(rely= 0.74,relx = 0.185,width = 185,height = 20)
        yscroll = ttk.Scrollbar(self.root, orient='vertical')
        self.tv3 = ttk.Treeview(self.combine_words, show="headings", columns=self.tv_tables3, height=30,
                                yscrollcommand=yscroll.set)
        self.tv3.place(rely=0.245, relx=0.02, width=570, height=150)
        yscroll.config(command=self.tv3.yview)
        for i in range(len(self.tv_tables3)):
            self.tv3.heading(column=self.tv_tables3[i], text=self.tv_tables3[i], anchor='center',
                               command=lambda: print(self.tv_tables3[i]))  # 设置除#0列以外的表头
            self.tv3.column(self.tv_tables3[i], width=25, anchor='center', stretch=True)
        if len(self.tree_view_data3):
            self.insert2treeview(tree_view_data=self.tree_view_data3, tv=self.tv3)
        self.tv3.bind("<<TreeviewSelect>>",self.on_tree_select7)


    def on_tree_select7(self,event):
        print("selected items:")
        self.tree_work_info3 = []
        for item in self.tv3.selection():
            self.tree_work_info3.append(self.tv3.item(item,'values')[0])
            # item_text = self.tv3.item(item,'values')[0]
        print(self.tree_work_info3)

    def fuc7_vis_menu(self):
        try:
            self.del_tree_items(self.tv3)
            self.tree_view_data3 = []
            menu_path = self.fuc07_1.get()
            for path in os.listdir(menu_path):
                self.tree_view_data3.append([path])
            self.insert2treeview(tree_view_data=self.tree_view_data3,tv = self.tv3)
        except:
            print('路径错误！')
    def fuc7_run(self):
        now = time.strftime("%Y-%m-%d %X",time.localtime())
        work_menu = self.tree_work_info3 # 需要处理的
        work_path = self.fuc07_1.get()
        combine_name = self.fuc07_2.get()
        word = wc.gencache.EnsureDispatch("Word.Application")
        new_document = word.Documents.Add()
        final_combine_path = os.path.join(work_path,combine_name)
        self.progressbar7['maximum'] = len(work_menu)
        self.progressbar7['value'] = 0
        for task in work_menu:
            path = os.path.join(work_path,task)
            new_document.Application.Selection.Range.InsertFile(path)
            new_document.SaveAs(final_combine_path)
            self.progressbar7['value'] += 1
            self.combine_words.update()
        new_document.Close()
        self.logger.insert('','end',value = [now,'合并word文档','正常'] )


    def show_fuc8(self):  #  批量生成pdf
        self.excel_related_f1_frame.place_forget()
        self.insert_font_f1_frame.place_forget()
        self.check_font_nums_frame.place_forget()
        self.auto_insert_name.place_forget()
        self.combine_words.place_forget()
        self.task_word_f1frame.place_forget()
        self.finall_check_scores.place_forget()
        self.auto_make_pdf.place(rely = 0.133,relx = 0.25)

        ttk.Label(self.auto_make_pdf,text = '使用说明:\n1.选择需要被转化成pdf的文档所在目录\n2.按照上述描述使用功能，等待进度条结束').place(rely = 0.82,relx = 0.02)
        ttk.Button(self.auto_make_pdf,text = '点击上传文档所在目录',width = 23,bootstyle="outline-dark",command = lambda :self.get_fuc_word_menu(self.fuc08_1)).place(relx = 0.02,rely = 0.02)
        ttk.Entry(self.auto_make_pdf,width = 31,textvariable = self.fuc08_1).place(relx = 0.325,rely = 0.02)
        ttk.Button(self.auto_make_pdf,text = '开始生成',command = lambda:self.pool.submit(self.run_fuc,self.fuc8_run),bootstyle="outline-dark",width = 10).place(relx = 0.02, rely = 0.72)
        ttk.Button(self.auto_make_pdf,text = '显示目录',command = self.fuc8_vis_menu,bootstyle="outline-dark").place(relx = 0.708,rely = 0.02)
        self.progressbar8.place(rely= 0.74,relx = 0.185,width = 185,height = 20)
        yscroll = ttk.Scrollbar(self.root, orient='vertical')
        self.tv4 = ttk.Treeview(self.auto_make_pdf, show="headings", columns=self.tv_tables4, height=30,
                                yscrollcommand=yscroll.set)
        self.tv4.place(rely=0.15, relx=0.02, width=570, height=185)
        yscroll.config(command=self.tv4.yview)
        for i in range(len(self.tv_tables4)):
            self.tv4.heading(column=self.tv_tables4[i], text=self.tv_tables4[i], anchor='center',)  # 设置除#0列以外的表头
            self.tv4.column(self.tv_tables4[i], width=25, anchor='center', stretch=True)

        if len(self.tree_work_info4):
            self.insert2treeview(tree_view_data=self.tree_view_data4, tv=self.tv4)
        self.tv4.bind("<<TreeviewSelect>>", self.on_tree_select8)


    def on_tree_select8(self,event):
        print("selected items:")
        self.tree_work_info4 = []

        for item in self.tv4.selection():
            self.tree_work_info4.append(self.tv4.item(item,'values')[0])
            # item_text = self.tv3.item(item,'values')[0]
        print(self.tree_work_info4)

    def fuc8_run(self):
        now = time.strftime("%Y-%m-%d %X",time.localtime())
        def createpdf(wordPath, pdfPath):
            word = gencache.EnsureDispatch('Word.Application')
            doc = word.Documents.Open(wordPath, ReadOnly=1)
            doc.ExportAsFixedFormat(pdfPath, constants.wdExportFormatPDF)
            word.Quit()
        try:
            self.logger.insert('end','开始生成pdf >>>>>>\n ')
            work_menu = self.tree_work_info4  # 需要处理的
            work_path = self.fuc08_1.get()
            self.progressbar8['maximum'] = len(work_menu)
            self.progressbar8['value'] = 0
            for file in work_menu:
                path = os.path.join(work_path,file)
                index =path.rindex('.')
                pdfpath = path[:index] + '.pdf'
                createpdf(path, pdfpath)
                self.progressbar8['value'] += 1
                self.auto_make_pdf.update()
                self.logger.insert('', 'end', value=[now, '批量生成pdf', '正常'])
        except FileNotFoundError:
                self.logger.insert('', 'end', value=[now, '批量生成pdf', '异常:检查路径或文件'])

    def fuc8_vis_menu(self):
        try:
            self.del_tree_items(self.tv4)
            self.tree_view_data4 = []
            menu_path = self.fuc08_1.get()
            for path in os.listdir(menu_path):
                self.tree_view_data4.append([path])
            self.insert2treeview(tree_view_data=self.tree_view_data4,tv = self.tv4)
        except:
            print('路径错误！')

    def show_fuc9(self):  #  成绩评定表批量插入签名
        self.excel_related_f1_frame.place_forget()
        self.insert_font_f1_frame.place_forget()
        self.check_font_nums_frame.place_forget()
        self.combine_words.place_forget()
        self.task_word_f1frame.place_forget()
        self.finall_check_scores.place_forget()
        self.auto_make_pdf.place_forget()
        self.auto_insert_name.place(rely = 0.133,relx = 0.25)
        p9 = threading.Thread(target = self.fuc9_run)
        ttk.Label(self.auto_insert_name,text = '使用说明:\n1.选择需要被插入签名的表所在目录(目录下只能存在此格式表,确保所有表都是关闭状态)\n2.选择您的身份并录入您的所要录入的签名\n3.点击开始，等待进度条结束便可完成').place(rely = 0.762,relx = 0.017)
        ttk.Button(self.auto_insert_name,text = '点击选择表所在目录',width = 23,bootstyle="outline-dark",command = lambda :self.get_fuc_word_menu(self.fuc09_1)).place(relx=0.017, rely=0.122)
        ttk.Entry(self.auto_insert_name,textvariable = self.fuc09_1,width = 31,state="readonly").place(relx = 0.325,rely = 0.122)
        ttk.Label(self.auto_insert_name,text = '选择您的身份：').place(rely = 0.25,relx = 0.02)
        ttk.Combobox(self.auto_insert_name,value = ['指导教师','评阅教师','答辩组长','答辩委员会负责人','答辩小组教师'],bootstyle="secondary",textvariable = self.fuc09_3).place( rely =  0.25,relx = 0.17,width = 125,height = 30)
        ttk.Label(self.auto_insert_name,text = '注意: 如果选择答辩记录表，请将身份选择为->答辩小组教师').place(relx = 0.02,rely = 0.345)
        ttk.Button(self.auto_insert_name,text = '录入签名图片',width = 10,bootstyle="outline-dark",command = lambda : self.get_fuc_label(self.fuc09_2)).place(relx=0.40, rely= 0.25)
        ttk.Button(self.auto_insert_name,text = '开始',command = lambda:self.pool.submit(self.run_fuc,self.fuc9_run),bootstyle="outline-dark",width = 10).place(relx = 0.02, rely = 0.65)
        # ttk.Button(self.auto_insert_name,text = '下载成绩评定表',bootstyle="dark-link",).place(rely = 0.02 ,relx = 0.72)
        # ttk.Button(self.auto_insert_name,text = '下载答辩记录表',bootstyle="dark-link",).place(rely = 0.122 ,relx = 0.72)
        self.progressbar9.place(rely= 0.665,relx = 0.185,width = 185,height = 20)
        ttk.Label(self.auto_insert_name,text = '选择需要处理的word类型：').place(relx = 0.02,rely = 0.02)
        ttk.Combobox(self.auto_insert_name,value = ['成绩评定表','答辩记录表'],bootstyle="secondary",textvariable = self.fuc09_4).place(relx = 0.325,rely = 0.02)

        # self.fuc9_text = ttk.ScrolledText(self.auto_insert_name,width = 80,state = 'normal',highlightcolor=style.colors.primary,highlightbackground = style.colors.border)
        # self.fuc9_text.place(rely = 0.39,relx = 0.02,height = 100)


    def fuc9_run(self):
            word_menu_path = self.fuc09_1.get()
            menu_list = os.listdir(word_menu_path)  # 处理的文档路径
            self.check_word_menu(menu_list,'docx')
            insert_pho = self.fuc09_2.get()  #  获取需要插入的签名
            user_id = self.fuc09_3.get()  #  用户的身份
            print('this is menu_list',menu_list)
            print('this is insert_pho',insert_pho)
            print('this is user_id',user_id)
            image = self.change_pic_size(insert_pho,(45, 35))
            new_path = insert_pho[:-4] + 'new.png'
            image.save(new_path)
            print('this is delete_path',new_path)
            word_class = self.fuc09_4.get()
            # self.fuc9_text.delete("1.0",'end')
            self.backups_word(word_menu_path)
            self.progressbar9['maximum'] = len(menu_list)
            self.progressbar9['value'] = 0
            for i in range(len(menu_list)):
                time.sleep(0.25)
                print('start ')
                path = word_menu_path + '\\'
                doc = Document(word_menu_path + '\\'+ menu_list[i])
                table = doc.tables[0]
                row = 0
                if user_id == '指导教师':
                    row = table.rows[2]
                elif user_id == '评阅教师':
                    row = table.rows[3]
                elif user_id == '答辩组长':
                    row = table.rows[4]
                elif user_id == '答辩委员会负责人':
                    row = table.rows[5]
                elif user_id == '答辩小组教师':
                    row = table.rows[16]
                cell = row.cells[0]
                if user_id == '答辩委员会负责人':
                    ph = cell.paragraphs[4]
                elif user_id == '答辩小组教师':
                    ph = cell.paragraphs[0]
                else:
                    ph = cell.paragraphs[6]
                run = ph.add_run()
                run.add_picture(new_path)
                path_name = path + '1'+ menu_list[i]
                doc.save(path_name)
                if os.path.exists(path + menu_list[i]):

                    os.remove(path + menu_list[i])
                os.rename(path_name,path + menu_list[i])
                # self.fuc9_text.insert("insert",'word:{}  身份:{}  日期:{} 处理成功! \n'.format(path,user_id,self.fuc09_4.get()))
                self.progressbar9['value'] += 1
                self.auto_insert_name.update()
            os.remove(new_path)
            print('插入完毕！！')

    @staticmethod
    def insert2treeview(tree_view_data,tv):
        for child in tree_view_data:
            tv.insert('', 'end', values=child)

    @staticmethod
    def check_word_menu(word_list,suffix):
        for i in range(len(word_list)-1):
            if not word_list[i].count(suffix):
                del word_list[i]


    @staticmethod
    def del_tree_items(tree):
        items = tree.get_children()
        for item in items:
            tree.delete(item)

    @staticmethod   # 备份原始数据
    def backups_word(path):
        backups_path = path + '\\' + 'raw_data'
        try:
            if not os.path.exists(backups_path):
                os.mkdir(backups_path)
                word_list = os.listdir(path)
                for i in word_list:
                    scr_path = path + '\\' + i
                    det_path = backups_path + '\\' + i
                    shutil.copyfile(scr_path,det_path)
            elif not len(os.listdir(backups_path)):
                word_list = os.listdir(path)
                for i in word_list:
                    scr_path = path + '\\' + i
                    det_path = backups_path + '\\'
                    shutil.copyfile(scr_path,det_path)
        except:
            pass

    @staticmethod
    def change_pic_size(pic,size):
        try:
            im = Image.open(pic)
            return im.resize(size)
        except:
            pass
    @staticmethod   # 保存路径
    def get_fuc_savepath(my_set):
        save_path = filedialog.askdirectory()
        my_set.set(save_path)
    @staticmethod   # 模板路径
    def get_fuc_model(my_set,title_name):
        model_path = filedialog.askopenfilename(title = title_name)
        my_set.set(model_path)
    @staticmethod   # 所要处理的表所在的路径
    def get_fuc_word_menu(my_set): # 获取成绩评定表所在目录
        save_path =filedialog.askdirectory(title = '请选择表所在目录')
        my_set.set(save_path)
    @staticmethod   # 获取签名所在的路径
    def get_fuc_label(my_set):
        label_path =  filedialog.askopenfilename(title = '选择签名图片')
        my_set.set(label_path )
    @staticmethod
    def get_fuc_excel(my_set):
        excel_path = filedialog.askopenfilename()
        my_set.set(excel_path)

    def doc2docx(self,file):   # 转doc 为 docx
        word = wc.Dispatch("Word.Application")  # 打开word应用程序
        doc = word.Documents.Open(file) #打开word文件
        doc.SaveAs("{}x".format(file), 12)#另存为后缀为".docx"的文件，其中参数12指docx文件
        doc.Close() #关闭原来word文件
        word.Quit()
        print('转化成功')

    @staticmethod
    def run_fuc(_fuc):
        _fuc()
        print('开始执行')

    def fuc6_excel_realte2combox(self):
        excel_path = filedialog.askopenfilename(title = '选择成绩评定表')
        self.fuc06_1.set(excel_path)
        if excel_path is not None:
            rets =  self.get_excel_para(pd.read_excel(excel_path))
            self.combobox6['values'] = rets


    @staticmethod
    def get_excel_para(df):  #  获取excel表的字段属性
        height, width = df.shape
        key_list = []
        for hgt in range(height):
            for col in range(width):
                if (pd.isna(df.iloc[hgt][col])):
                    key_list.clear()
                    break
                key_list.append(df.iloc[hgt][col])
            if (len(key_list) - width == 0):
                break
        return key_list
if __name__ == '__main__':

    root = tk.Tk()
    style = ttk.Style()
    style.configure('left.Label',font =('黑体',18),background = '#666699',)
    style.configure('left_frame.TFrame',background = '#99CCFF')
    MainPage(master = root)
    style.configure('left_button.TButton',font =('黑体',12),background = '#99CCFF',bootstyle = 'info')
    root.mainloop()